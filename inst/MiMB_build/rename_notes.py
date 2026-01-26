#!/usr/bin/env python3
"""Rename Note 4.X, Figure 5.X, and Table 5.X references in DOCX file."""

import re
from docx import Document


def rename_references(input_path, output_path):
    """Replace 'Note 4.X' -> 'Note X', 'Figure 5.X' -> 'Figure X', 'Table 5.X' -> 'Table X'."""
    doc = Document(input_path)

    # Pattern for Note heading number in its own run (e.g., "4.1" or "4.12")
    note_heading_pattern = re.compile(r'^4\.(\d+)$')

    # Pattern for "4.X" or "5.X" in its own run (handles split references)
    section_number_pattern = re.compile(r'^([45])\.(\d+)$')

    # Pattern for "Figure 5.X:" or "Table 5.X:" at start of caption (often in single run)
    figure_caption_pattern = re.compile(r'^(Figure\s+)5\.(\d+)(:.*)?$')
    table_caption_pattern = re.compile(r'^(Table\s+)5\.(\d+)(:.*)?$')

    def process_paragraph(paragraph):
        """Process all runs in a paragraph."""
        runs = paragraph.runs

        for i, run in enumerate(runs):
            text = run.text
            if not text:
                continue

            # Check if this run is exactly a Note heading number like "4.1" or "4.12"
            # (typically the first run in a heading paragraph)
            if i == 0 and note_heading_pattern.match(text):
                run.text = note_heading_pattern.sub(r'\1', text)
                continue

            # Check if this run starts with "Figure 5.X:" (caption format)
            if figure_caption_pattern.match(text):
                run.text = figure_caption_pattern.sub(r'\g<1>\2\3', text)
                continue

            # Check if this run starts with "Table 5.X:" (caption format)
            if table_caption_pattern.match(text):
                run.text = table_caption_pattern.sub(r'\g<1>\2\3', text)
                continue

            # Check if this run is "4.X" or "5.X" following "Note", "Figure", or "Table"
            match = section_number_pattern.match(text)
            if match:
                prefix_num = match.group(1)  # "4" or "5"
                sub_num = match.group(2)     # The subsection number

                # Look backwards to find the preceding word
                prev_text = ""
                for j in range(i - 1, -1, -1):
                    prev_text = (runs[j].text or "") + prev_text
                    if prev_text.strip():
                        break

                # Check context: "Note" before "4.X", or "Figure"/"Table" before "5.X"
                if prefix_num == "4" and re.search(r'Note\s*$', prev_text):
                    run.text = sub_num
                elif prefix_num == "5" and re.search(r'(Figure|Table)\s*$', prev_text):
                    run.text = sub_num

    # Process all paragraphs
    for para in doc.paragraphs:
        process_paragraph(para)

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_paragraph(para)

    # Process headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header:
                for para in header.paragraphs:
                    process_paragraph(para)
        for footer in [section.footer, section.first_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    process_paragraph(para)

    doc.save(output_path)
    print(f"Saved renumbered document to: {output_path}")


if __name__ == "__main__":
    input_file = "MiMBIntegratedPTM.docx"
    output_file = "MiMBIntegratedPTM_renumbered.docx"
    rename_references(input_file, output_file)
